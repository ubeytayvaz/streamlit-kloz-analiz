import streamlit as st
import PyPDF2
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import io
import re
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes

# Aranacak klozlar, Türkçe karşılıkları ve ilgili anahtar kelimeler
CLAUSE_DEFINITIONS = {
    "Total Asbestos Exclusion Clause": {
        "tr": "Total Asbest İstisnası Klozu",
        "keywords": ["asbestos", "asbest", "asbesto"]
    },
    "CL380 Institute Cyber Attack Exclusion": {
        "tr": "CL380 Siber Saldırı İstisnası Enstitü Klozu",
        "keywords": ["cyber", "siber", "CL380", "hacker", "bilgisayar virüsü", "computer virus", "cyber attack", "siber saldırı"]
    },
    "LMA5394 Communicable Disease Exclusion": {
        "tr": "LMA5394 Bulaşıcı Hastalık İstisnası Klozu",
        "keywords": ["communicable disease", "bulaşıcı hastalık", "LMA5394", "salgın", "epidemic", "pandemic", "pandemi"]
    },
    "NMA 2738 Claims Control Clause": {
        "tr": "NMA 2738 Hasar Kontrol Klozu",
        "keywords": ["claims control", "hasar kontrol", "NMA 2738", "claims"]
    },
    "LMA3100 Sanction Limitation and Exclusion Clause": {
        "tr": "LMA3100 Yaptırım Sınırlama ve İstisna Klozu",
        "keywords": ["sanction", "yaptırım", "LMA3100", "ambargo", "embargo", "limitation", "sınırlama"]
    },
    "Contingent Business Interruption": {
        "tr": "Dolaylı Kar Kaybı / Koşullu İş Durması",
        "keywords": [
            "contingent business interruption", "dolaylı kar kaybı", "tedarikçi riski", 
            "supplier risk", "müşteri riski", "customer risk", "erişimin engellenmesi", 
            "prevention of access", "denial of access", "public authorities", "kamu otoriteleri"
        ]
    }
}

def extract_text_from_pdf(file_content):
    """PDF'ten metin ve sayfa numaralarını çıkarır."""
    pages_content = []
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                pages_content.append({'page': i + 1, 'content': page_text})
        return pages_content
    except Exception:
        return []

def extract_text_with_ocr(file_content, is_pdf=True):
    """OCR kullanarak taranmış dosyalardan metin ve sayfa numaralarını çıkarır."""
    pages_content = []
    try:
        if is_pdf:
            images = convert_from_bytes(file_content)
        else:
            images = [Image.open(io.BytesIO(file_content))]
        
        st.info(f"OCR işlemi başlatıldı. {len(images)} sayfa/resim metne dönüştürülüyor...")
        lang_config = 'eng+tur'
        progress_bar = st.progress(0)
        
        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image, lang=lang_config)
            pages_content.append({'page': i + 1, 'content': text})
            progress_bar.progress((i + 1) / len(images))
        progress_bar.empty()
        return pages_content
    except Exception as e:
        st.error(f"OCR işlemi sırasında bir hata oluştu: {e}")
        return None

def extract_text_from_docx(file_content):
    """DOCX'ten metin çıkarır ve tek sayfa olarak kabul eder."""
    try:
        doc = Document(io.BytesIO(file_content))
        text = "\n".join([para.text for para in doc.paragraphs])
        return [{'page': 1, 'content': text}]
    except Exception as e:
        st.error(f"DOCX dosyası okunurken bir hata oluştu: {e}")
        return None

def analyze_document_text(pages_content, clause_definitions):
    """Metin içinde klozları ve anahtar kelimeleri sayfa bazında arar."""
    found_items = []
    processed_texts = set()

    for page_data in pages_content:
        page_num = page_data['page']
        text = page_data['content']
        sentences = re.split(r'(?<=[.!?])\s+', text)

        for en_clause, details in clause_definitions.items():
            # Tam eşleşmeleri ara
            for clause_name in [en_clause, details.get("tr", "")]:
                if not clause_name: continue
                pattern = re.compile(re.escape(clause_name), re.IGNORECASE)
                for match in pattern.finditer(text):
                    found_text = match.group(0).strip()
                    if f"{found_text.lower()}_{page_num}" not in processed_texts:
                        found_items.append({"clause": en_clause, "found_text": found_text, "reason": "Tam eşleşme bulundu.", "page": page_num})
                        processed_texts.add(f"{found_text.lower()}_{page_num}")
            
            # Anahtar kelimeleri ara
            for keyword in details.get("keywords", []):
                if not keyword: continue
                pattern = re.compile(r'\b' + re.escape(keyword) + r'\b', re.IGNORECASE)
                for sentence in sentences:
                    if pattern.search(sentence):
                        clean_sentence = sentence.strip()
                        if clean_sentence and f"{clean_sentence.lower()}_{page_num}" not in processed_texts:
                            found_items.append({"clause": en_clause, "found_text": clean_sentence, "reason": f"İlgili anahtar kelime: '{keyword}'", "page": page_num})
                            processed_texts.add(f"{clean_sentence.lower()}_{page_num}")
    return found_items

def create_highlighted_html(full_text, found_items):
    highlighted_text = full_text
    sorted_items = sorted(found_items, key=lambda x: len(x['found_text']), reverse=True)
    for item in sorted_items:
        pattern = re.compile(re.escape(item['found_text']), re.IGNORECASE)
        highlighted_text = pattern.sub(lambda match: f"<mark>{match.group(0)}</mark>" if "<mark>" not in match.group(0) else match.group(0), highlighted_text)
    return highlighted_text

def highlight_text_in_pdf(file_content, texts_to_highlight):
    try:
        pdf_doc = fitz.open(stream=file_content, filetype="pdf")
        for page in pdf_doc:
            for text in texts_to_highlight:
                text_instances = page.search_for(text)
                for inst in text_instances:
                    highlight = page.add_highlight_annot(inst)
                    highlight.update()
        output_buffer = io.BytesIO()
        pdf_doc.save(output_buffer)
        pdf_doc.close()
        return output_buffer.getvalue()
    except Exception as e:
        st.error(f"PDF işaretlenirken hata: {e}")
        return None

def highlight_text_in_docx(file_content, texts_to_highlight):
    try:
        doc = Document(io.BytesIO(file_content))
        for para in doc.paragraphs:
            if any(text.lower() in para.text.lower() for text in texts_to_highlight):
                for run in para.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        output_buffer = io.BytesIO()
        doc.save(output_buffer)
        return output_buffer.getvalue()
    except Exception as e:
        st.error(f"DOCX işaretlenirken hata: {e}")
        return None

# --- Streamlit Arayüzü ---

st.set_page_config(page_title="OCR Destekli Poliçe Kloz Analiz Aracı", layout="wide")

st.title("🤖 OCR Destekli Poliçe Kloz Analiz Aracı")
st.markdown("Bu araç, dijital veya **taranmış** poliçe dosyalarınızı (.pdf, .docx, .png, .jpg) analiz ederek belirli klozları ve ilgili anahtar kelimeleri tespit eder.")

with st.expander("Analiz Edilen Klozlar ve Anahtar Kelimeleri Gör"):
    for clause, details in CLAUSE_DEFINITIONS.items():
        st.info(f"**{clause}**")
        st.markdown(f"**Anahtar Kelimeler:** `{'`, `'.join(details['keywords'])}`")

uploaded_file = st.file_uploader(
    "Lütfen analiz edilecek dosyayı seçin (.pdf, .docx, .png, .jpg).",
    type=["pdf", "docx", "png", "jpg", "jpeg"]
)

if uploaded_file is not None:
    if 'processed_file' not in st.session_state or st.session_state.processed_file != uploaded_file.name:
        st.session_state.processed_file = uploaded_file.name
        
        file_content = uploaded_file.read()
        file_type = uploaded_file.type
        
        with st.spinner(f"'{uploaded_file.name}' içeriği okunuyor..."):
            pages_content = []
            if file_type == "application/pdf":
                pages_content = extract_text_from_pdf(file_content)
                if not pages_content or len("".join(p['content'] for p in pages_content).strip()) < 100:
                    st.warning("Bu bir taranmış PDF gibi görünüyor. Metin okuma (OCR) işlemi başlatılıyor...")
                    pages_content = extract_text_with_ocr(file_content, is_pdf=True)
            elif file_type.startswith('image/'):
                pages_content = extract_text_with_ocr(file_content, is_pdf=False)
            else: # docx
                pages_content = extract_text_from_docx(file_content)
        
        st.session_state.pages_content = pages_content
        st.session_state.file_content = file_content
    
    if 'pages_content' in st.session_state and st.session_state.pages_content:
        file_name = uploaded_file.name
        pages_content = st.session_state.pages_content
        file_content = st.session_state.file_content
        full_document_text = "\n".join([page['content'] for page in pages_content])

        with st.spinner("Belge analiz ediliyor..."):
            found_items = analyze_document_text(pages_content, CLAUSE_DEFINITIONS)
        
        st.success(f"'{file_name}' analizi tamamlandı. {len(found_items)} potansiyel bulgu tespit edildi.")
        st.markdown("---")
        
        st.header("🔍 Analiz Sonuçları")

        col1, col2 = st.columns([1, 2])

        with col1:
            if found_items:
                st.subheader("✅ Tespit Edilen Bulgular")
                
                grouped_results = {}
                for item in found_items:
                    clause = item['clause']
                    if clause not in grouped_results:
                        grouped_results[clause] = []
                    grouped_results[clause].append(item)

                for clause, items in grouped_results.items():
                    with st.expander(f"İlgili Kloz: {clause} ({len(items)} bulgu)"):
                        for item in sorted(items, key=lambda x: x['page']):
                            st.markdown(f"**Sayfa:** {item['page']}")
                            st.markdown(f"**Sebep:** {item['reason']}")
                            st.info(f"*{item['found_text']}*")
                
                st.markdown("---")
                st.subheader("⬇️ İşaretli Dosyayı İndir")
                
                if file_name.lower().endswith(('.pdf', '.docx')):
                    texts_to_highlight = [item['found_text'] for item in found_items]
                    highlighted_file_bytes = None
                    
                    with st.spinner("Dosya işaretleniyor..."):
                        if file_name.lower().endswith('.pdf'):
                            highlighted_file_bytes = highlight_text_in_pdf(file_content, texts_to_highlight)
                        else:
                            highlighted_file_bytes = highlight_text_in_docx(file_content, texts_to_highlight)

                    if highlighted_file_bytes:
                        st.download_button(label="İndirmek için tıklayın", data=highlighted_file_bytes, file_name=f"isaretli_{file_name}", mime=uploaded_file.type)
                else:
                    st.info("İşaretli dosya indirme özelliği şimdilik sadece PDF ve DOCX dosyaları için geçerlidir.")
            else:
                st.warning("Listelenen klozlar veya ilgili anahtar kelimeler bu belgede bulunamadı.")
        
        with col2:
            st.subheader("📑 Belge İçeriği Önizlemesi")
            highlighted_html = create_highlighted_html(full_document_text, found_items)
            st.markdown(f'<div style="background-color:#f0f2f6; border: 1px solid #ddd; border-radius: 5px; padding: 15px; height: 600px; overflow-y: scroll;">{highlighted_html.replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
    
    elif 'pages_content' in st.session_state:
         st.error("Dosyadan metin okunamadı. Lütfen dosyanın bozuk veya şifreli olmadığını kontrol edin.")

